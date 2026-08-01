from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Nasze-Wesele-dokument-techniczno-biznesowy.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_keep(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)


def style_run(run, bold=None, italic=None, color=None, size=None) -> None:
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def add_para(doc, text="", style=None, after=6, before=0, bold=False, italic=False, color=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        r = p.add_run(text)
        style_run(r, bold=bold, italic=italic, color=color)
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def _find_abstract_num_id(doc: Document, style_id: str) -> str:
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        for lvl in abstract.findall(qn("w:lvl")):
            p_style = lvl.find(qn("w:pStyle"))
            if p_style is not None and p_style.get(qn("w:val")) == style_id:
                return abstract.get(qn("w:abstractNumId"))
    return "7"


def _new_numbering_id(doc: Document, style_id: str = "ListNumber") -> str:
    numbering = doc.part.numbering_part.element
    existing = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) and num.get(qn("w:numId")).isdigit()
    ]
    num_id = str((max(existing) if existing else 0) + 1)
    abstract_num_id = _find_abstract_num_id(doc, style_id)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def _apply_num_id(paragraph, num_id: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), num_id)


def add_numbers(doc, items: list[str]) -> None:
    num_id = _new_numbering_id(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        _apply_num_id(p, num_id)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    set_paragraph_keep(p)
    return p


def add_callout(doc, title: str, body: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, [9360])
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    style_run(r, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    p2.add_run(body)
    add_para(doc, "", after=4)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, widths)
    repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        style_run(r, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.add_run(value)
    set_table_width(table, widths)
    add_para(doc, "", after=4)


def add_status_matrix(doc, rows: list[list[str]]) -> None:
    add_table(doc, ["Obszar", "Stan obecny", "Implikacja dla analizy/testow"], rows, [1900, 2850, 4610])


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Nasze Wesele | dokument techniczno-biznesowy"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(9)
    hp.runs[0].font.color.rgb = GRAY

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Dokument roboczy dla analitykow i testerow systemowo-biznesowych")
    fp.runs[0].font.size = Pt(9)
    fp.runs[0].font.color.rgb = GRAY


def build_doc() -> None:
    doc = Document()
    configure_doc(doc)

    # Cover / masthead
    add_para(doc, "DOKUMENT TECHNICZNO-BIZNESOWY", after=2, bold=True, color=BLUE)
    title = add_para(doc, "Nasze Wesele", after=4, bold=True, color=NAVY)
    title.runs[0].font.size = Pt(28)
    subtitle = add_para(doc, "Specyfikacja produktu, zakres analizy, ryzyka oraz pakiet testow systemowo-biznesowych", after=16, color=GRAY)
    subtitle.runs[0].font.size = Pt(13)

    add_table(
        doc,
        ["Metadana", "Wartosc"],
        [
            ["Odbiorcy", "Analitycy biznesowi, analitycy systemowi, testerzy systemowi, testerzy UAT, product owner, osoby odpowiedzialne za wdrozenie."],
            ["System", "Nasze Wesele - aplikacja webowa SaaS/MVP dla obslugi stron weselnych, panelu pary, panelu operatora i procesu sprzedazy."],
            ["Repozytorium lokalne", "D:\\wedding"],
            ["Wersja dokumentu", "1.0"],
            ["Data opracowania", date.today().isoformat()],
            ["Podstawa", "Przeglad kodu, README, dokumentacji, migracji Supabase, konfiguracji Next.js i uruchomionych walidacji lint/build."],
        ],
        [2100, 7260],
    )

    add_callout(
        doc,
        "Najwazniejszy wniosek",
        "Projekt jest funkcjonalnym MVP z szeroka warstwa prezentacyjna i sprzedazowa. Nadaje sie do demonstracji, analizy wymagan i pierwszych manualnych wdrozen, ale przed publiczna produkcja wymaga domkniecia autoryzacji, zapisu danych tenantowych w Supabase, realnego uploadu do Storage oraz zabezpieczenia API operatorskiego.",
    )

    add_heading(doc, "Spis tresci logicznej", 1)
    add_numbers(
        doc,
        [
            "Cel dokumentu i sposob uzycia przez zespoly.",
            "Opis produktu, aktorzy, wartosc biznesowa i model sprzedazy.",
            "Zakres funkcjonalny systemu oraz mapowanie modulow.",
            "Architektura, dane, integracje, konfiguracja i bezpieczenstwo.",
            "Procesy biznesowe i reguly decyzyjne.",
            "Wymagania funkcjonalne, niefunkcjonalne i ograniczenia MVP.",
            "Strategia testow, scenariusze E2E, macierz pokrycia i kryteria akceptacji.",
            "Ryzyka, luki produkcyjne, rekomendacje i roadmapa.",
        ],
    )

    add_heading(doc, "1. Cel i zakres dokumentu", 1)
    add_para(doc, "Dokument laczy perspektywe biznesowa, systemowa i testowa. Ma byc praktycznym materialem roboczym dla osob, ktore beda doprecyzowywac wymagania, projektowac przypadki testowe, prowadzic UAT oraz oceniac gotowosc systemu do wdrozenia komercyjnego.")
    add_heading(doc, "1.1 Jak korzystac z dokumentu", 2)
    add_bullets(
        doc,
        [
            "Analityk biznesowy powinien zaczac od sekcji produktu, aktorow, procesow i regul biznesowych.",
            "Analityk systemowy powinien przejsc przez architekture, dane, API, konfiguracje i ograniczenia techniczne.",
            "Tester systemowy powinien uzyc macierzy modulow, scenariuszy E2E, checklist regresji oraz kryteriow akceptacji.",
            "Tester biznesowy/UAT powinien weryfikowac zgodnosc z procesem sprzedazy, onboardingiem pary, obsluga goscia i wymaganiami prawnymi.",
            "Product owner powinien traktowac sekcje luk i roadmapy jako backlog gotowosci produkcyjnej.",
        ],
    )

    add_heading(doc, "1.2 Definicje", 2)
    add_table(
        doc,
        ["Pojecie", "Znaczenie"],
        [
            ["Tenant", "Pojedyncze wesele/instancja w systemie, identyfikowana przez slug i docelowo przez wedding_id."],
            ["Slug", "Unikalny fragment adresu URL, np. anna-michal, uzywany w /w/[slug] oraz /app/[slug]."],
            ["Para", "Klient systemu: para mloda albo osoba organizujaca wesele."],
            ["Gosc", "Osoba korzystajaca z publicznej strony wesela, RSVP, wyszukiwarki stolika, uploadu i galerii."],
            ["Operator", "Osoba sprzedajaca i konfigurujaca instancje, korzystajaca ze /start, /sprzedaz, /zamowienie i /admin."],
            ["MVP", "Zakres minimalny, ktory dziala pokazowo i czesciowo lokalnie, ale nie ma jeszcze pelnej automatyzacji produkcyjnej."],
            ["RLS", "Row Level Security w Supabase, czyli polityki ograniczajace dostep do rekordow po roli i wedding_id."],
        ],
        [1800, 7560],
    )

    add_heading(doc, "2. Streszczenie wykonawcze", 1)
    add_para(doc, "Nasze Wesele to aplikacja webowa dla obslugi wesela jako cyfrowego centrum informacji. Produkt ma obslugiwac gosci przez publiczna strone i QR, pare przez panel administracyjny, a operatora przez panel sprzedazy i kreator instancji. System jest przygotowany koncepcyjnie pod multi-tenant SaaS, jednak aktualny stan implementacji miesza dane demo/localStorage z czesciowo przygotowanym Supabase.")
    add_status_matrix(
        doc,
        [
            ["Frontend i UX", "Rozbudowane strony, komponenty goscia, panel admina, oferta, materialy i motywy.", "Testy powinny objac nawigacje, responsywnosc, czytelnosc, formularze i zgodnosc tresci z pakietami."],
            ["Build i jakosc kodu", "Lint oraz build produkcyjny przechodza bez bledow.", "Mozna rozpoczac formalne testy systemowe na lokalnym lub testowym srodowisku."],
            ["Dane demo", "Znaczna czesc panelu i danych dziala w localStorage.", "Testy musza rozroznic zachowanie demo od produkcyjnego zapisu w bazie."],
            ["Supabase", "Sa migracje, seed, klient admin i API sprzedazowe.", "Wymagane testy integracyjne po skonfigurowaniu projektu Supabase i RLS."],
            ["Bezpieczenstwo", "Istnieje Basic Auth dla ekranow operatorskich po ustawieniu hasla, security headers i rate limit w middleware.", "API operatorskie i panel tenantowy wymagaja dodatkowego auth przed produkcja."],
            ["Upload", "UI waliduje liczbe, typ i rozmiar plikow, ale nie zapisuje realnie do storage.", "Nie traktowac uploadu jako gotowej funkcji produkcyjnej; testowac jako mock do czasu implementacji Storage."],
        ],
    )

    add_heading(doc, "3. Produkt i model biznesowy", 1)
    add_heading(doc, "3.1 Propozycja wartosci", 2)
    add_para(doc, "Produkt sprzedaje cyfrowe centrum wesela pod linkiem i kodem QR. Gosc ma szybko znalezc plan dnia, miejsce przy stole, lokalizacje, FAQ, galerie i formularze. Para ma dostac panel organizacyjny, w ktorym moze zarzadzac danymi wesela, goscmi, stolikami, dokumentami, QR i komunikatami. Operator ma moc szybko sprzedac, utworzyc instancje i obsluzyc klienta recznie przed pelnym self-service.")
    add_heading(doc, "3.2 Pakiety komercyjne", 2)
    add_table(
        doc,
        ["Pakiet", "Cena", "Limity", "Kluczowy zakres"],
        [
            ["Start", "199 zl", "0 GB, 0 min wideo, 12 mies.", "Strona weselna, harmonogram, lokalizacje, FAQ, RSVP, QR do strony."],
            ["Wesele Live", "349 zl", "10 GB, 30 min wideo, 12 mies.", "Start plus upload zdjec/wideo, galeria, pokaz slajdow, ksiega gosci, QR do druku."],
            ["Organizer Pro", "599 zl", "25 GB, 90 min wideo, 18 mies.", "Live plus plan stolow, mapa sali, planner, umowy, zaliczki, dokumenty."],
            ["Concierge", "1200 zl+", "50 GB, 180 min wideo, 24 mies.", "Pro plus import gosci, konfiguracja, QR, wsparcie przed weselem, priorytet."],
        ],
        [1500, 1200, 2050, 4610],
    )
    add_heading(doc, "3.3 Etapy monetyzacji", 2)
    add_numbers(
        doc,
        [
            "Etap manualny: operator sprzedaje pakiet, pobiera platnosc recznie i tworzy instancje w /zamowienie.",
            "Etap operacyjny: operator zarzadza leadami i instancjami w /sprzedaz oraz konfiguruje klienta.",
            "Etap self-service: klient placi online, system tworzy konto, instancje, subskrypcje, onboarding i wysyla email.",
            "Etap partnerski: produkt jest sprzedawany przez sale weselne, wedding plannerow, fotografow i DJ-ow.",
        ],
    )

    add_heading(doc, "4. Aktorzy i uprawnienia", 1)
    add_table(
        doc,
        ["Aktor", "Cele", "Dostep / ekrany", "Ryzyka testowe"],
        [
            ["Gosc", "Uzyskac informacje, znalezc stolik, odpowiedziec RSVP, dodac zdjecia, wpisac sie do ksiegi.", "/, /upload, /gallery, /rsvp, /guestbook, /music oraz docelowo /w/[slug].", "Brak logowania; trzeba testowac prywatnosc, kody dostepu, walidacje i zachowanie na mobile."],
            ["Para", "Skonfigurowac wesele, gosci, stoliki, QR, dokumenty, galerie i komunikaty.", "/app/[slug], /app/[slug]/panel, obecnie tez demo /admin.", "Docelowo wymaga Supabase Auth i RLS; obecnie dane tenantowe sa lokalne."],
            ["Operator", "Sprzedaz, tworzenie instancji, obsluga leadow, materialy.", "/start, /oferta, /zamowienie, /sprzedaz, /materialy.", "Basic Auth jest opcjonalny; API wymaga domkniecia zabezpieczen."],
            ["Superadmin", "Zarzadzanie leadami, instancjami, planami, aktywacjami i wsparciem.", "Docelowo rola w Supabase app_superadmins.", "Nalezy testowac separacje danych, role i brak dostepu publicznego."],
        ],
        [1350, 2600, 2500, 2910],
    )

    add_heading(doc, "5. Zakres funkcjonalny systemu", 1)
    add_heading(doc, "5.1 Mapa modulow", 2)
    add_table(
        doc,
        ["Modul", "Sciezki / pliki", "Opis biznesowy", "Status"],
        [
            ["Publiczna strona wesela", "/, /w/[slug], komponenty guest", "Hero, plan dnia, informacje, wyszukiwarka goscia, lokalizacje, FAQ, menu, kontakt.", "Demo rozbudowane; tenant /w/[slug] nie pobiera jeszcze pelnych danych z Supabase."],
            ["RSVP", "/rsvp, rsvp-form", "Potwierdzenie obecnosci i informacje od goscia.", "UI demo; do integracji z tabela guests/households."],
            ["Ksiega gosci", "/guestbook, guestbook-form", "Zbieranie zyczen z moderacja.", "Model bazy przewiduje zatwierdzanie; UI wymaga testow zapisu."],
            ["Ankieta muzyczna", "/music, song-request-form", "Prosby muzyczne od gosci.", "UI funkcjonalny w warstwie demo."],
            ["Upload i galeria", "/upload, /gallery, /slideshow", "Dodawanie zdjec/wideo, galeria i pokaz slajdow.", "Upload symulowany; brak realnego Storage."],
            ["Panel admina pary", "/admin, /app/[slug]/panel, AdminPanel", "Dashboard, goscie, stoliki, QR, galeria, motyw, dokumenty, planner.", "Bardzo szeroki UI, ale oparty lokalnie o localStorage."],
            ["Sprzedaz", "/oferta, /zamowienie, /sprzedaz", "Landing, leady, kreator instancji, panel operatora.", "Czesciowo podpiete pod API/Supabase, fallback lokalny."],
            ["Materialy", "/materialy, /materialy/qr, exports", "PDF-y, QR do druku, materialy sprzedazowe.", "Gotowe artefakty w repo."],
        ],
        [1500, 2100, 3400, 2360],
    )

    add_heading(doc, "5.2 Zakres poza MVP lub niedomkniety", 2)
    add_bullets(
        doc,
        [
            "Pelny self-service platnosci Stripe/Przelewy24 i webhooki.",
            "Automatyczne tworzenie konta wlasciciela oraz magic link do panelu.",
            "Pelna autoryzacja panelu /app/[slug] przez Supabase Auth.",
            "Trwaly zapis konfiguracji panelu pary po wedding_id.",
            "Realny upload do Supabase Storage z limitami pakietu, moderacja i retencja.",
            "Superadmin z rolami, audytem i pelnym zabezpieczeniem endpointow operatorskich.",
            "Monitoring produkcyjny, alerty kosztowe, backup bazy i storage.",
        ],
    )

    add_heading(doc, "6. Architektura systemu", 1)
    add_heading(doc, "6.1 Stack technologiczny", 2)
    add_table(
        doc,
        ["Warstwa", "Technologia", "Uwagi"],
        [
            ["Frontend / SSR", "Next.js 16.2.7, React 19.2.4, TypeScript", "App Router, strony statyczne i dynamiczne, middleware."],
            ["Style", "Tailwind CSS 4, global CSS", "Mobile-first, komponenty produktowe i operatorskie."],
            ["Ikony/UI", "lucide-react, sonner, dnd-kit", "Ikony, powiadomienia, potencjalny drag-and-drop w panelu."],
            ["Formularze/walidacja", "react-hook-form, zod", "Walidacja formularzy i schematow."],
            ["Backend danych", "Supabase Postgres/Auth/Storage", "Migracje i RLS przygotowane, integracja czesciowa."],
            ["Eksporty", "html2canvas, jsPDF, jszip, xlsx", "Materialy QR/PDF/XLSX i eksport danych."],
        ],
        [1900, 2750, 4710],
    )

    add_heading(doc, "6.2 Model uruchomienia", 2)
    add_numbers(
        doc,
        [
            "Lokalnie: npm install, npm run dev, aplikacja na http://localhost:3000.",
            "Walidacja: npm run lint i npm run build.",
            "Supabase: uruchomienie migracji i seed w projekcie Supabase.",
            "Vercel: ustawienie zmiennych srodowiskowych, domeny, security i monitoringu.",
            "Produkcja: konieczne ustawienie INTERNAL_TOOLS_PASSWORD i docelowo auth/role-based access.",
        ],
    )

    add_heading(doc, "6.3 Zmienne srodowiskowe", 2)
    add_table(
        doc,
        ["Zmienna", "Przeznaczenie", "Wymagalnosc"],
        [
            ["NEXT_PUBLIC_SUPABASE_URL", "Adres projektu Supabase.", "Wymagana dla integracji z baza."],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Klucz publiczny Supabase dla klienta.", "Wymagana dla klienta Supabase."],
            ["SUPABASE_SERVICE_ROLE_KEY", "Klucz serwisowy do operacji API po stronie serwera.", "Wymagana dla API sprzedazowych; nie moze trafic do klienta."],
            ["NEXT_PUBLIC_SITE_URL", "Adres bazowy aplikacji.", "Wymagana dla linkow publicznych i QR."],
            ["INTERNAL_TOOLS_PASSWORD", "Haslo Basic Auth dla /admin, /start, /sprzedaz.", "Krytyczne przed publicznym wdrozeniem."],
            ["UPLOAD_MAX_FILE_SIZE_MB, UPLOAD_MAX_FILES", "Limity uploadu konfiguracyjne.", "Do spiecia z UI i backendem."],
            ["PAYMENT_PROVIDER, STRIPE_*, PRZELEWY24_*", "Konfiguracja platnosci.", "Na razie opcjonalna/manualna."],
            ["RESEND_API_KEY, SENTRY_DSN", "Maile i monitoring bledow.", "Wymagane dla produkcji."],
        ],
        [2600, 4300, 2460],
    )

    add_heading(doc, "7. Dane i model domenowy", 1)
    add_heading(doc, "7.1 Glowne encje", 2)
    add_table(
        doc,
        ["Encja", "Opis", "Najwazniejsze relacje/testy"],
        [
            ["weddings", "Konfiguracja wesela: slug, para, data, miejsca, publikacja, plan, limity.", "Test unikalnosci slug, publikacji, wygasania, limitow i widocznosci publicznej."],
            ["wedding_admins", "Powiazanie uzytkownika Supabase Auth z weselem.", "Test RLS: admin widzi tylko swoje wedding_id."],
            ["households", "Zaproszenia rodzinne/grupowe z kodem.", "Test kodu zaproszenia i przypisania gosci."],
            ["guests", "Lista gosci, RSVP, menu, muzyka, stolik.", "Test walidacji RSVP, preferencji i wyszukiwarki."],
            ["seating_tables", "Stoliki, nazwa, pojemnosc, mapa sali.", "Test pojemnosci, przypisan, mapy i konfliktow."],
            ["schedule_items", "Elementy harmonogramu dnia.", "Test sortowania, widocznosci i edycji."],
            ["faq_items", "FAQ widoczne publicznie.", "Test publikacji i kolejnosci."],
            ["announcements", "Komunikaty dla gosci.", "Test published_at, aktywnosci i priorytetu."],
            ["guestbook_entries", "Wpisy ksiegi gosci z moderacja.", "Test dodania publicznego i zatwierdzenia przez admina."],
            ["sales_leads", "Leady z formularza oferty.", "Test walidacji email, planu, slug i statusow."],
            ["wedding_subscriptions", "Pakiet, status, kwota, daty dostepu.", "Test statusow platnosci, wygasania i blokad."],
            ["wedding_onboarding_steps", "Kroki konfiguracji po zakupie.", "Test kompletacji i widocznosci w panelu."],
        ],
        [2200, 4300, 2860],
    )

    add_heading(doc, "7.2 Reguly danych tenantowych", 2)
    add_bullets(
        doc,
        [
            "Kazdy rekord biznesowy powinien byc przypisany do wedding_id, gdy dotyczy konkretnego wesela.",
            "Slug powinien miec 3-64 znaki, tylko male litery, cyfry i myslniki, bez znakow diakrytycznych.",
            "Publiczne odczyty powinny dotyczyc tylko opublikowanych wesel i zatwierdzonych tresci.",
            "Panel pary nie moze czytac ani modyfikowac danych innego wedding_id.",
            "Wszystkie operacje service role musza byc wykonywane tylko po stronie serwera.",
            "Dane osobowe gosci, RSVP i pliki wymagaja retencji, eksportu/usuniecia oraz jasnych zgod.",
        ],
    )

    add_heading(doc, "8. API i integracje", 1)
    add_table(
        doc,
        ["Endpoint", "Metoda", "Cel", "Stan / ryzyko"],
        [
            ["/api/ping", "GET", "Health check.", "Do prostego monitoringu dostepnosci."],
            ["/api/sales/leads", "POST", "Zapis leada do sales_leads.", "Wymaga service role; publiczny formularz powinien miec dodatkowy anti-spam/captcha."],
            ["/api/sales/checkout", "POST", "Utworzenie weddings, subscription i onboarding steps.", "Aktualnie model manualny; trzeba testowac rollback/kompensacje przy bledzie czesciowym."],
            ["/api/sales/instances", "GET", "Lista instancji dla panelu sprzedazy.", "Krytyczne: endpoint wymaga auth/superadmin przed produkcja."],
        ],
        [2200, 900, 2900, 3360],
    )
    add_heading(doc, "8.1 Integracje docelowe", 2)
    add_bullets(
        doc,
        [
            "Supabase Auth: logowanie pary i operatora, role, magic links, przypisanie do wedding_admins.",
            "Supabase Storage: zdjecia, wideo, dokumenty, eksporty, podpisane linki.",
            "Stripe lub Przelewy24: checkout, webhook, status subskrypcji, fakturowanie lub integracja ksiegowa.",
            "Resend: maile onboardingowe, potwierdzenia, linki do panelu.",
            "Sentry: monitoring bledow frontend/backend.",
            "Plausible/Umami: analityka ruchu landingow i konwersji.",
        ],
    )

    add_heading(doc, "9. Bezpieczenstwo, prywatnosc i zgodnosc", 1)
    add_heading(doc, "9.1 Mechanizmy obecne", 2)
    add_bullets(
        doc,
        [
            "Security headers w next.config.ts: X-Frame-Options, nosniff, Referrer-Policy, CSP, Permissions-Policy.",
            "Middleware z rate limitingiem dla /api oraz ekranow operatorskich.",
            "Opcjonalny Basic Auth dla /admin, /start i /sprzedaz po ustawieniu INTERNAL_TOOLS_PASSWORD.",
            "Supabase RLS w migracjach dla tabel domenowych.",
            "Oddzielenie klienta admin Supabase do uzycia server-side z service role.",
        ],
    )
    add_heading(doc, "9.2 Luki przed produkcja", 2)
    add_table(
        doc,
        ["Luka", "Skutek", "Rekomendacja testowo-analityczna"],
        [
            ["Brak pelnego auth dla /app/[slug]", "Nie ma produkcyjnej separacji panelu pary.", "Zaprojektowac i przetestowac Supabase Auth + wedding_admins + RLS."],
            ["GET /api/sales/instances bez jawnego auth", "Ryzyko wycieku listy klientow po skonfigurowaniu Supabase.", "Dodac superadmin check lub zabezpieczenie middleware/API."],
            ["Publiczny lead form bez captcha", "Ryzyko spamu i kosztow.", "Dodac Turnstile/captcha i rate limit per formularz."],
            ["Upload tylko po stronie klienta", "Walidacja moze byc ominieta.", "Walidowac typ, rozmiar, limit pakietu i status platnosci na backendzie."],
            ["Brak audytu zmian", "Trudniej odtworzyc kto zmienil dane.", "Dodac created_by/updated_by i log operacji dla paneli."],
            ["CSP z unsafe-inline/unsafe-eval", "Szersza powierzchnia XSS.", "Zweryfikowac wymagania Next/Turbopack i zaostrzyc CSP w produkcji."],
        ],
        [2500, 3000, 3860],
    )

    add_heading(doc, "10. Procesy biznesowe", 1)
    add_heading(doc, "10.1 Sprzedaz manualna i utworzenie instancji", 2)
    add_numbers(
        doc,
        [
            "Klient trafia na /oferta i wybiera pakiet lub wysyla zapytanie.",
            "Lead zapisuje sie w Supabase albo lokalnie w fallbacku przegladarki.",
            "Operator kontaktuje sie z klientem, ustala zakres i pobiera platnosc manualnie.",
            "Operator przechodzi do /zamowienie i wpisuje dane pary, date, email, telefon, pakiet i slug.",
            "System generuje publicPath /w/[slug], adminPath /app/[slug] i storagePrefix weddings/[slug].",
            "Przy skonfigurowanym Supabase API tworzy rekordy weddings, wedding_subscriptions i wedding_onboarding_steps.",
            "Operator przekazuje parze link do panelu i materialy QR.",
        ],
    )
    add_heading(doc, "10.2 Onboarding pary", 2)
    add_numbers(
        doc,
        [
            "Para otrzymuje link do panelu.",
            "Para uzupelnia dane wesela, miejsca, godziny, komunikaty i kontakt.",
            "Para dodaje lub importuje gosci oraz przypisuje ich do gospodarstw/zaproszen.",
            "Para uklada stoliki i plan sali.",
            "Para wybiera motyw, grafike, tresci FAQ i ustawienia widocznosci.",
            "System generuje QR do strony, uploadu, galerii, RSVP i wyszukiwarki stolika.",
            "Przed weselem para publikuje strone i testuje dostep goscia.",
        ],
    )
    add_heading(doc, "10.3 Dzien wesela", 2)
    add_bullets(
        doc,
        [
            "Gosc skanuje QR i przechodzi do strony wesela.",
            "Gosc sprawdza harmonogram, lokalizacje, menu, FAQ i swoje miejsce przy stole.",
            "Gosc dodaje zdjecia lub krotkie wideo, jesli pakiet i status platnosci na to pozwalaja.",
            "Para lub operator moderuje tresci, galerie i ksiege gosci.",
            "Pokaz slajdow moze byc wyswietlany na sali.",
        ],
    )

    add_heading(doc, "11. Wymagania funkcjonalne", 1)
    functional_rows = [
        ["FR-001", "System pozwala wyswietlic landing publiczny z informacjami o weselu.", "MVP", "Gosc"],
        ["FR-002", "System pozwala wyszukac goscia i jego stolik.", "MVP", "Gosc"],
        ["FR-003", "System pozwala wyswietlic harmonogram, lokalizacje, FAQ, menu i kontakty.", "MVP", "Gosc"],
        ["FR-004", "System pozwala dodac RSVP i dane preferencji goscia.", "MVP/docelowo baza", "Gosc"],
        ["FR-005", "System pozwala dodac wpis do ksiegi gosci z moderacja.", "MVP/docelowo baza", "Gosc/Para"],
        ["FR-006", "System pozwala zglosic prosbe muzyczna.", "MVP", "Gosc"],
        ["FR-007", "System pozwala wybrac pliki do uploadu i waliduje typ/rozmiar/liczbe.", "MVP UI", "Gosc"],
        ["FR-008", "System zapisuje pliki w Storage i respektuje limity pakietu.", "Do zrobienia", "Gosc/System"],
        ["FR-009", "Operator moze przyjac lead z formularza oferty.", "MVP", "Operator"],
        ["FR-010", "Operator moze utworzyc instancje wesela po zakupie.", "MVP", "Operator"],
        ["FR-011", "Operator moze przegladac instancje w panelu sprzedazy.", "MVP", "Operator"],
        ["FR-012", "Para moze zarzadzac danymi wesela, goscmi, stolikami, QR, dokumentami i motywem.", "MVP localStorage", "Para"],
        ["FR-013", "Para widzi tylko swoje wesele i dane po wedding_id.", "Do zrobienia", "Para"],
        ["FR-014", "Superadmin moze zarzadzac leadami, instancjami i aktywacja.", "Czesciowo", "Superadmin"],
        ["FR-015", "System obsluguje platnosc online i webhook aktywacji.", "Do zrobienia", "System"],
    ]
    add_table(doc, ["ID", "Wymaganie", "Status", "Aktor"], functional_rows, [900, 5200, 1800, 1460])

    add_heading(doc, "12. Wymagania niefunkcjonalne", 1)
    add_table(
        doc,
        ["Kategoria", "Wymaganie", "Miernik akceptacji"],
        [
            ["Wydajnosc", "Strona goscia musi ladowac szybko na telefonie i slabym internecie.", "Lighthouse mobile do ustalenia; brak krytycznych blokad renderu."],
            ["Dostepnosc", "Kluczowe scenariusze goscia musza dzialac na mobile.", "Chrome/Safari mobile, viewporty 360-430 px, brak nakladania tekstu."],
            ["Bezpieczenstwo", "Panel i API operatorskie niedostepne publicznie.", "Brak dostepu bez auth; test bez hasla zwraca 401/403."],
            ["Prywatnosc", "Dane gosci i pliki sa separowane per wesele.", "Test RLS i storage policy dla min. dwoch tenantow."],
            ["Niezawodnosc", "Bledy Supabase/API nie psuja calego UI.", "Komunikaty bledow i fallbacki kontrolowane."],
            ["Skalowalnosc kosztow", "Upload wideo nie moze generowac niekontrolowanych kosztow.", "Limity per pakiet egzekwowane backendowo."],
            ["Utrzymywalnosc", "Build i lint przechodza na CI.", "npm run lint i npm run build bez bledow."],
            ["Zgodnosc prawna", "System ma zgody, regulamin, retencje i procedury usuniecia.", "Dokumenty prawne zatwierdzone i linkowane w UI."],
        ],
        [1800, 4400, 3160],
    )

    add_heading(doc, "13. Reguly biznesowe", 1)
    add_table(
        doc,
        ["ID", "Regula", "Test / warunek brzegowy"],
        [
            ["BR-001", "Slug wesela jest normalizowany do malych liter, cyfr i myslnikow, maksymalnie 64 znaki.", "Wprowadzic polskie znaki, spacje, znaki specjalne, dlugi tekst."],
            ["BR-002", "Pakiet okresla limit storage, minut wideo i czas dostepu.", "Porownac Start/Live/Pro/Concierge i blokady po przekroczeniu."],
            ["BR-003", "Wesele nieoplacone ma status platnosc oczekujaca i nie powinno byc w pelni publiczne.", "Utworzyc instancje manual i paid; sprawdzic is_published/status."],
            ["BR-004", "Publiczny gosc moze zobaczyc tylko opublikowane dane i zatwierdzone wpisy.", "Dodac wpis niezatwierdzony, sprawdzic brak widocznosci publicznej."],
            ["BR-005", "Wpisy ksiegi gosci wymagaja moderacji.", "Nowy wpis powinien miec is_approved=false."],
            ["BR-006", "Upload musi byc limitowany liczba plikow, typem, rozmiarem i pakietem.", "Test 0 plikow, 31 plikow, zly MIME, >25 MB, pakiet Start."],
            ["BR-007", "Operator powinien miec dostep tylko po zabezpieczeniu narzedzi wewnetrznych.", "Bez INTERNAL_TOOLS_PASSWORD system nie powinien byc wystawiany publicznie."],
            ["BR-008", "Dane panelu pary musza byc izolowane per wedding_id.", "Dwa konta, dwa wesela, proba odczytu obcego wedding_id."],
        ],
        [900, 5550, 2910],
    )

    add_heading(doc, "14. Strategia testow", 1)
    add_heading(doc, "14.1 Poziomy testow", 2)
    add_table(
        doc,
        ["Poziom", "Cel", "Zakres"],
        [
            ["Smoke", "Potwierdzic, ze aplikacja startuje i kluczowe strony odpowiadaja.", "/, /oferta, /zamowienie, /sprzedaz, /admin, /w/test, /app/test/panel, API ping."],
            ["Systemowe", "Zweryfikowac funkcje koncowe i integracje.", "Formularze, tenant routing, panel, API Supabase, fallbacki."],
            ["Systemowo-biznesowe", "Zweryfikowac zgodnosc z procesem sprzedazy, onboardingu i dnia wesela.", "Lead -> zamowienie -> instancja -> panel -> QR -> strona goscia."],
            ["Regresyjne", "Utrzymac stabilnosc po zmianach.", "Nawigacja, formularze, build, auth, RLS, upload, mobile."],
            ["UAT", "Potwierdzic akceptacje przez operatora i pare.", "Praca na realistycznym weselu demo z minimum 30 goscmi i 5 stolikami."],
            ["Bezpieczenstwa", "Sprawdzic role, separacje danych i brak wyciekow.", "Endpointy, middleware, RLS, service role, Storage policies."],
        ],
        [1850, 3200, 4310],
    )

    add_heading(doc, "14.2 Dane testowe", 2)
    add_bullets(
        doc,
        [
            "Minimum dwa tenanty: anna-michal oraz kasia-piotr.",
            "Minimum dwa konta pary przypisane do roznych wedding_id.",
            "Minimum jeden superadmin i jeden zwykly uzytkownik bez przypisania.",
            "Goscie: para, rodzina z jednym kodem zaproszenia, dzieci, dieta specjalna, osoba bez stolika, osoba niepotwierdzona.",
            "Stoliki: pelny stolik, wolny stolik, stolik VIP, konflikt pojemnosci.",
            "Pliki: jpg/png/webp/heic/mp4/mov, plik za duzy, plik o zlym MIME, 31 plikow naraz.",
            "Pakiety: Start bez uploadu, Live z 10 GB, Pro z dokumentami, Concierge z importem.",
        ],
    )

    add_heading(doc, "15. Scenariusze E2E", 1)
    e2e_rows = [
        ["E2E-001", "Lead z oferty", "Uzytkownik wypelnia formularz /oferta.", "Lead zapisany w Supabase albo fallback localStorage, komunikat sukcesu."],
        ["E2E-002", "Manualne zamowienie", "Operator tworzy instancje w /zamowienie.", "Powstaje /w/[slug], /app/[slug], subscription i onboarding steps."],
        ["E2E-003", "Panel sprzedazy", "Operator otwiera /sprzedaz.", "Widzi leady i instancje, moze rozroznic statusy."],
        ["E2E-004", "Konfiguracja pary", "Para przechodzi do /app/[slug]/panel i edytuje dane.", "Zmiany widoczne w panelu danego sluga; docelowo zapis w Supabase."],
        ["E2E-005", "Publiczna strona goscia", "Gosc otwiera /w/[slug] i przechodzi do strony.", "Widoczne informacje weselne i CTA; docelowo dane z bazy."],
        ["E2E-006", "RSVP goscia", "Gosc wypelnia RSVP.", "Walidacja pol, zapis statusu i brak duplikatow."],
        ["E2E-007", "Wyszukiwarka stolika", "Gosc wpisuje imie/nazwisko.", "System pokazuje stolik, osoby przy stoliku i komunikaty prywatne."],
        ["E2E-008", "Upload zdjec", "Gosc wybiera poprawne i niepoprawne pliki.", "System odrzuca naruszenia i przyjmuje poprawne pliki; produkcyjnie zapisuje do Storage."],
        ["E2E-009", "Moderacja ksiegi", "Gosc dodaje wpis, para zatwierdza.", "Wpis nie jest publiczny przed zatwierdzeniem, potem pojawia sie w widoku."],
        ["E2E-010", "Separacja tenantow", "Admin A probuje odczytac dane admina B.", "Dostep zablokowany przez RLS/auth."],
    ]
    add_table(doc, ["ID", "Scenariusz", "Warunek startowy / akcja", "Oczekiwany rezultat"], e2e_rows, [900, 1800, 3350, 3310])

    add_heading(doc, "16. Checklisty testowe", 1)
    add_heading(doc, "16.1 Smoke test", 2)
    add_bullets(
        doc,
        [
            "npm run lint konczy sie bez bledow.",
            "npm run build konczy sie bez bledow.",
            "Strona / renderuje sie na desktop i mobile.",
            "Panel /admin renderuje sie lokalnie.",
            "Oferta /oferta renderuje pakiety i formularz.",
            "Kreator /zamowienie pozwala przejsc przez dane testowe.",
            "Panel /sprzedaz pokazuje fallback albo dane z API.",
            "Dynamiczne sciezki /w/test-slug oraz /app/test-slug/panel nie zwracaja 500.",
            "API /api/ping odpowiada poprawnie.",
        ],
    )
    add_heading(doc, "16.2 Regresja funkcjonalna", 2)
    add_bullets(
        doc,
        [
            "Nawigacja miedzy publicznymi modulami dziala bez martwych linkow.",
            "Formularze waliduja wymagane pola, email, daty i limity dlugosci.",
            "Slug generuje sie przewidywalnie z polskich znakow i spacji.",
            "Pakiety w UI zgadzaja sie z konfiguracja commercial-config.",
            "Zmiany localStorage nie mieszaja danych miedzy slugami.",
            "Rate limit zwraca 429 po przekroczeniu progu.",
            "Basic Auth blokuje /admin, /start i /sprzedaz po ustawieniu hasla.",
            "Endpointy API nie ujawniaja danych przy braku wymaganej konfiguracji/auth.",
            "Na mobile teksty nie wychodza poza przyciski i karty.",
        ],
    )
    add_heading(doc, "16.3 Testy negatywne", 2)
    add_bullets(
        doc,
        [
            "Pusty formularz leadu.",
            "Niepoprawny email w leadzie i checkout.",
            "Slug z samymi znakami specjalnymi.",
            "Brak SUPABASE_SERVICE_ROLE_KEY przy wywolaniu API wymagajacego bazy.",
            "Blad Supabase przy drugim insercie checkoutu.",
            "Upload bez pliku, za duzy plik, zly MIME, zbyt wiele plikow.",
            "Dostep do /sprzedaz bez Basic Auth po ustawieniu hasla.",
            "Proba pobrania danych innego tenant_id przez konto pary.",
            "Nieopublikowane wesele widoczne publicznie.",
        ],
    )

    add_heading(doc, "17. Macierz pokrycia testowego", 1)
    add_table(
        doc,
        ["Obszar", "Priorytet", "Typ testu", "Minimalne pokrycie"],
        [
            ["Sprzedaz i leady", "Wysoki", "Systemowe + biznesowe", "Lead poprawny, lead bledny, fallback, zapis Supabase."],
            ["Checkout / zamowienie", "Wysoki", "E2E + integracyjne", "Manual/paid, pakiety, slug, subscription, onboarding."],
            ["Auth i role", "Krytyczny", "Bezpieczenstwa", "Operator, superadmin, para, brak roli, dwa tenanty."],
            ["Panel pary", "Wysoki", "Systemowe + UAT", "Edycja danych, goscie, stoliki, QR, motyw, dokumenty."],
            ["Strona goscia", "Wysoki", "Systemowe + mobile", "Plan, lokalizacje, stolik, FAQ, galerie, CTA."],
            ["Upload", "Krytyczny po wdrozeniu Storage", "Integracyjne + negatywne", "Typy, rozmiary, pakiety, status platnosci, moderacja."],
            ["RLS / dane", "Krytyczny", "Integracyjne", "Widocznosc publiczna i izolacja wedding_id."],
            ["Platnosci", "Krytyczny po implementacji", "E2E + webhook", "Sukces, anulowanie, duplikat webhooka, aktywacja."],
            ["Responsywnosc", "Sredni/wysoki", "UI", "360, 390, 430, 768, 1440 px."],
            ["Dokumenty prawne", "Sredni", "Biznesowe", "Linki, tresci zgody, retencja, usuniecie danych."],
        ],
        [2500, 1400, 2300, 3160],
    )

    add_heading(doc, "18. Kryteria akceptacji wydania", 1)
    add_heading(doc, "18.1 MVP demonstracyjne", 2)
    add_bullets(
        doc,
        [
            "Build i lint przechodza.",
            "Wszystkie ekrany z README sa dostepne lokalnie.",
            "Sprzedazowy flow demo dziala z fallbackiem localStorage.",
            "Panel admina umozliwia pokazanie podstawowej konfiguracji wesela.",
            "Materialy PDF/QR sa dostepne w exports i /materialy.",
            "Znane ograniczenia sa jawnie opisane dla klienta lub zespolu.",
        ],
    )
    add_heading(doc, "18.2 Pierwsze platne wdrozenie manualne", 2)
    add_bullets(
        doc,
        [
            "Supabase jest skonfigurowany, migracje i seed przeszly poprawnie.",
            "INTERNAL_TOOLS_PASSWORD jest ustawiony na produkcji.",
            "Operator ma procedury tworzenia instancji i przekazania linkow.",
            "Leady i checkout zapisuja sie w bazie.",
            "Panel klienta nie jest publiczny bez logowania albo jest swiadomie obslugiwany manualnie bez danych wrażliwych.",
            "Upload produkcyjny jest albo wylaczony, albo realnie podpiety do Storage z limitami backendowymi.",
            "Regulamin, polityka prywatnosci, zgody i retencja sa gotowe.",
        ],
    )
    add_heading(doc, "18.3 Publiczny SaaS self-service", 2)
    add_bullets(
        doc,
        [
            "Platnosci online i webhooki sa odporne na duplikaty i bledy.",
            "Konto pary tworzy sie automatycznie i jest przypisane do wedding_admins.",
            "RLS potwierdzone testami dla minimum dwoch tenantow.",
            "Wszystkie dane panelu pary sa w Supabase, nie tylko localStorage.",
            "Storage ma polityki dostepu, podpisane linki, limity i monitoring kosztow.",
            "Superadmin ma role, audyt i zabezpieczone API.",
            "Monitoring, backup, alerty i procedury incydentowe sa wdrozone.",
        ],
    )

    add_heading(doc, "19. Ryzyka i rekomendacje", 1)
    add_table(
        doc,
        ["Ryzyko", "Wplyw", "Prawdopodobienstwo", "Rekomendacja"],
        [
            ["Wyciek danych operatorskich przez API", "Wysoki", "Srednie", "Zabezpieczyc /api/sales/instances rola superadmin i testami negatywnymi."],
            ["Brak separacji danych pary", "Wysoki", "Wysokie do czasu auth", "Wdrozyc Supabase Auth i RLS w panelu /app/[slug]."],
            ["Upload generuje koszty", "Wysoki", "Srednie", "Limity backendowe, quota per pakiet, blokada po wygasnieciu."],
            ["Za szeroki zakres MVP", "Sredni", "Wysokie", "Priorytetyzowac flow sprzedaz -> instancja -> panel -> publiczna strona."],
            ["Tanie pakiety generuja support", "Sredni", "Wysokie", "Na start sprzedawac Concierge/Pro i ograniczac personalizacje."],
            ["Niejasny status prawny zdjec", "Wysoki", "Srednie", "Zatwierdzic zgody, retencje, usuniecie danych i administratora danych."],
            ["Brak monitoringu produkcji", "Sredni", "Srednie", "Dodac Sentry, logi, alerty i health checks."],
        ],
        [3050, 1350, 1850, 3110],
    )

    add_heading(doc, "20. Backlog rekomendowany", 1)
    add_table(
        doc,
        ["Priorytet", "Zadanie", "Uzasadnienie", "Kryterium gotowosci"],
        [
            ["P0", "Zabezpieczyc API operatorskie i /api/sales/instances.", "Bez tego dane klientow moga byc narazone.", "Brak dostepu bez superadmin/auth; testy negatywne."],
            ["P0", "Wdrozyc Supabase Auth dla /app/[slug].", "Panel pary musi byc prywatny.", "Admin widzi tylko swoje wedding_id."],
            ["P0", "Przeniesc dane panelu z localStorage do Supabase.", "Trwalosc i multi-tenant wymagaja bazy.", "Edycja danych utrzymuje sie po odswiezeniu i na innym urzadzeniu."],
            ["P0", "Podpiac upload do Supabase Storage.", "Kluczowa wartosc pakietu Live/Pro.", "Pliki sa zapisane, limitowane, moderowane i widoczne w galerii."],
            ["P1", "Dodac platnosci i webhooki.", "Self-service i skalowanie sprzedazy.", "Platnosc aktywuje pakiet bez recznej interwencji."],
            ["P1", "Dodac onboarding mailowy.", "Zmniejszenie supportu.", "Para dostaje link, instrukcje i checklist."],
            ["P1", "Dodac audit log i monitoring.", "Operacje na danych osobowych i wsparcie produkcji.", "Logi zmian i bledow widoczne dla operatora."],
            ["P2", "Automatyczne eksporty i backup.", "Zaufanie klientow i bezpieczenstwo danych.", "Backup bazy/storage i eksport wesela."],
        ],
        [1000, 3000, 3200, 2160],
    )

    add_heading(doc, "21. Zalecenia dla analitykow", 1)
    add_bullets(
        doc,
        [
            "Oddzielic wymagania demo od wymagan produkcyjnych i oznaczac je w backlogu.",
            "Opisac kompletne BPMN lub user journey dla sprzedazy manualnej, onboardingu pary i dnia wesela.",
            "Zdefiniowac polityke prywatnosci dla danych gosci, zdjec, wideo i dokumentow.",
            "Doprecyzowac role: operator, superadmin, wlasciciel wesela, wspoladministrator, gosc.",
            "Ustalic, ktore funkcje sa w pakietach Start/Live/Pro/Concierge i co system robi po przekroczeniu limitu.",
            "Przygotowac definicje statusow: lead, instancja, subskrypcja, publikacja, platnosc, onboarding.",
            "Zaprojektowac komunikaty bledow i fallbackow dla klienta nietechnicznego.",
        ],
    )

    add_heading(doc, "22. Zalecenia dla testerow", 1)
    add_bullets(
        doc,
        [
            "Nie zakladac, ze pozytywny komunikat w UI oznacza trwaly zapis w bazie; weryfikowac persistence.",
            "Kazdy test tenantowy wykonywac na minimum dwoch slugach i dwoch kontach.",
            "Testowac mobile jako glowny kanal goscia, nie jako dodatek.",
            "W scenariuszach uploadu testowac rowniez backend i Storage policy, gdy zostana wdrozone.",
            "W testach bezpieczenstwa sprawdzac endpointy bezposrednio, nie tylko ekrany.",
            "Przy testach biznesowych porownywac UI z pakietami cenowymi i dokumentami prawnymi.",
            "Raportowac defekty z etykieta: demo-limitacja, blad produkcyjny, luka bezpieczenstwa, niespojnosc biznesowa.",
        ],
    )

    add_heading(doc, "23. Zalaczniki", 1)
    add_heading(doc, "23.1 Lista glownych sciezek aplikacji", 2)
    add_table(
        doc,
        ["Sciezka", "Przeznaczenie"],
        [
            ["/", "Demo publicznej strony wesela."],
            ["/admin", "Demo panelu administracyjnego."],
            ["/oferta", "Landing sprzedazowy i formularz leadow."],
            ["/start", "Kokpit produktu/operatora."],
            ["/zamowienie", "Kreator instancji po zakupie."],
            ["/sprzedaz", "Panel leadow i instancji."],
            ["/materialy", "Materialy sprzedazowe."],
            ["/materialy/qr", "QR do druku."],
            ["/motywy", "Motywy demo."],
            ["/demo", "Demo dla par."],
            ["/w/[slug]", "Publiczna sciezka tenantowa wesela."],
            ["/app/[slug]", "Docelowy panel konkretnej instancji."],
            ["/app/[slug]/panel", "Panel konfiguracji konkretnego wesela."],
            ["/upload", "Upload zdjec/wideo od gosci."],
            ["/gallery", "Galeria gosci."],
            ["/slideshow", "Pokaz slajdow."],
            ["/rsvp", "RSVP."],
            ["/guestbook", "Ksiega gosci."],
            ["/music", "Prosby muzyczne."],
        ],
        [2400, 6960],
    )
    add_heading(doc, "23.2 Komendy walidacyjne", 2)
    add_table(
        doc,
        ["Komenda", "Cel", "Oczekiwany wynik"],
        [
            ["npm install", "Instalacja zaleznosci.", "Brak bledow instalacji."],
            ["npm run dev", "Start lokalny.", "Aplikacja dziala na localhost:3000."],
            ["npm run lint", "Statyczna walidacja kodu.", "Brak bledow ESLint."],
            ["npm run build", "Build produkcyjny.", "Kompilacja, TypeScript i generowanie stron bez bledow."],
            ["supabase db reset", "Migracje i seed lokalnego Supabase.", "Demo anna-michal/nasze-wesele dostepne w bazie."],
        ],
        [2200, 3400, 3760],
    )
    add_heading(doc, "23.3 Decyzje otwarte", 2)
    add_bullets(
        doc,
        [
            "Czy pierwsze wdrozenia maja blokowac upload do czasu Storage, czy dopuszczamy upload manualny poza systemem?",
            "Czy Basic Auth wystarczy tymczasowo na produkcji, czy od razu wdrazamy Supabase Auth dla operatorow?",
            "Czy para moze samodzielnie publikowac strone, czy publikacja jest zatwierdzana przez operatora?",
            "Jak dlugo po weselu utrzymujemy dane, zdjecia i panel?",
            "Jakie sa zasady zwrotow, reklamacji i zmiany pakietu?",
            "Czy produkt sprzedajemy B2C parom, czy priorytetem jest B2B przez sale/plannerow?",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc()
    print(OUT)
